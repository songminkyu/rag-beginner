"""
Document Loader Module
다양한 형식의 문서를 로딩하는 모듈
"""

import os
import csv
import json
import logging
from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
from dataclasses import dataclass

# 외부 라이브러리 (optional imports)
try:
    import PyPDF2
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """문서 표준화 클래스"""
    content: str
    metadata: Dict[str, Any]
    source: str
    doc_id: Optional[str] = None
    
    def __post_init__(self):
        if self.doc_id is None:
            # 파일명을 기반으로 ID 생성
            self.doc_id = os.path.basename(self.source)


class DocumentLoader(ABC):
    """문서 로더 기본 클래스"""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        self.config = config or {}
        self.encoding = self.config.get("encoding", "utf-8")
    
    @abstractmethod
    def load(self, source: Union[str, Path]) -> List[Document]:
        """문서 로딩"""
        pass
    
    @abstractmethod
    def supports_file_type(self, file_path: str) -> bool:
        """파일 타입 지원 여부 확인"""
        pass
    
    def load_directory(
        self, 
        directory: Union[str, Path],
        recursive: bool = True,
        file_pattern: Optional[str] = None
    ) -> List[Document]:
        """디렉토리 내 모든 지원 파일 로딩"""
        
        directory = Path(directory)
        if not directory.exists():
            raise FileNotFoundError(f"디렉토리를 찾을 수 없습니다: {directory}")
        
        documents = []
        
        if recursive:
            pattern = "**/*" if file_pattern is None else f"**/{file_pattern}"
        else:
            pattern = "*" if file_pattern is None else file_pattern
        
        for file_path in directory.glob(pattern):
            if file_path.is_file() and self.supports_file_type(str(file_path)):
                try:
                    docs = self.load(file_path)
                    documents.extend(docs)
                    logger.info(f"로딩 완료: {file_path} ({len(docs)}개 문서)")
                except Exception as e:
                    logger.error(f"파일 로딩 실패: {file_path}, 오류: {e}")
        
        return documents
    
    def validate_file(self, file_path: Union[str, Path]) -> bool:
        """파일 유효성 검사"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"파일이 존재하지 않습니다: {file_path}")
            return False
        
        if not file_path.is_file():
            logger.error(f"파일이 아닙니다: {file_path}")
            return False
        
        if file_path.stat().st_size == 0:
            logger.warning(f"빈 파일입니다: {file_path}")
            return False
        
        return True


class TextLoader(DocumentLoader):
    """텍스트 파일 로더"""
    
    def load(self, source: Union[str, Path]) -> List[Document]:
        """텍스트 파일 로딩"""
        
        if not self.validate_file(source):
            return []
        
        try:
            with open(source, 'r', encoding=self.encoding, errors='ignore') as f:
                content = f.read()
            
            metadata = {
                "file_type": "text",
                "file_size": os.path.getsize(source),
                "encoding": self.encoding,
                "loader": "TextLoader"
            }
            
            return [Document(
                content=content,
                metadata=metadata,
                source=str(source)
            )]
            
        except Exception as e:
            logger.error(f"텍스트 파일 로딩 오류: {source}, {e}")
            return []
    
    def supports_file_type(self, file_path: str) -> bool:
        """텍스트 파일 타입 지원 여부"""
        text_extensions = {'.txt', '.md', '.rst', '.log', '.py', '.js', '.html', '.css'}
        return Path(file_path).suffix.lower() in text_extensions


class PDFLoader(DocumentLoader):
    """PDF 파일 로더"""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__(config)
        
        if not HAS_PDF:
            raise ImportError("PDF 로딩을 위해 PyPDF2와 pdfplumber를 설치해주세요: pip install PyPDF2 pdfplumber")
        
        self.use_pdfplumber = self.config.get("use_pdfplumber", True)
        self.extract_images = self.config.get("extract_images", False)
    
    def load(self, source: Union[str, Path]) -> List[Document]:
        """PDF 파일 로딩"""
        
        if not self.validate_file(source):
            return []
        
        if self.use_pdfplumber:
            return self._load_with_pdfplumber(source)
        else:
            return self._load_with_pypdf2(source)
    
    def _load_with_pdfplumber(self, source: Union[str, Path]) -> List[Document]:
        """pdfplumber를 사용한 PDF 로딩"""
        
        documents = []
        
        try:
            with pdfplumber.open(source) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    
                    if text and text.strip():
                        metadata = {
                            "file_type": "pdf",
                            "page_number": page_num,
                            "total_pages": len(pdf.pages),
                            "file_size": os.path.getsize(source),
                            "loader": "PDFLoader-pdfplumber"
                        }
                        
                        documents.append(Document(
                            content=text,
                            metadata=metadata,
                            source=str(source),
                            doc_id=f"{os.path.basename(source)}-page-{page_num}"
                        ))
            
            logger.info(f"PDF 로딩 완료: {source} ({len(documents)}페이지)")
            return documents
            
        except Exception as e:
            logger.error(f"PDF 로딩 오류 (pdfplumber): {source}, {e}")
            return []
    
    def _load_with_pypdf2(self, source: Union[str, Path]) -> List[Document]:
        """PyPDF2를 사용한 PDF 로딩"""
        
        documents = []
        
        try:
            with open(source, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(reader.pages, 1):
                    text = page.extract_text()
                    
                    if text and text.strip():
                        metadata = {
                            "file_type": "pdf",
                            "page_number": page_num,
                            "total_pages": len(reader.pages),
                            "file_size": os.path.getsize(source),
                            "loader": "PDFLoader-PyPDF2"
                        }
                        
                        documents.append(Document(
                            content=text,
                            metadata=metadata,
                            source=str(source),
                            doc_id=f"{os.path.basename(source)}-page-{page_num}"
                        ))
            
            logger.info(f"PDF 로딩 완료: {source} ({len(documents)}페이지)")
            return documents
            
        except Exception as e:
            logger.error(f"PDF 로딩 오류 (PyPDF2): {source}, {e}")
            return []
    
    def supports_file_type(self, file_path: str) -> bool:
        """PDF 파일 타입 지원 여부"""
        return Path(file_path).suffix.lower() == '.pdf'


class DocxLoader(DocumentLoader):
    """DOCX 파일 로더"""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__(config)
        
        if not HAS_DOCX:
            raise ImportError("DOCX 로딩을 위해 python-docx를 설치해주세요: pip install python-docx")
    
    def load(self, source: Union[str, Path]) -> List[Document]:
        """DOCX 파일 로딩"""
        
        if not self.validate_file(source):
            return []
        
        try:
            doc = DocxDocument(source)
            
            # 모든 문단 텍스트 추출
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            content = '\n'.join(paragraphs)
            
            metadata = {
                "file_type": "docx",
                "paragraph_count": len(paragraphs),
                "file_size": os.path.getsize(source),
                "loader": "DocxLoader"
            }
            
            return [Document(
                content=content,
                metadata=metadata,
                source=str(source)
            )]
            
        except Exception as e:
            logger.error(f"DOCX 파일 로딩 오류: {source}, {e}")
            return []
    
    def supports_file_type(self, file_path: str) -> bool:
        """DOCX 파일 타입 지원 여부"""
        return Path(file_path).suffix.lower() == '.docx'


class CSVLoader(DocumentLoader):
    """CSV 파일 로더"""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__(config)
        
        self.delimiter = self.config.get("delimiter", ",")
        self.text_column = self.config.get("text_column", None)
        self.combine_columns = self.config.get("combine_columns", True)
    
    def load(self, source: Union[str, Path]) -> List[Document]:
        """CSV 파일 로딩"""
        
        if not self.validate_file(source):
            return []
        
        documents = []
        
        try:
            if HAS_PANDAS:
                return self._load_with_pandas(source)
            else:
                return self._load_with_csv(source)
            
        except Exception as e:
            logger.error(f"CSV 파일 로딩 오류: {source}, {e}")
            return []
    
    def _load_with_pandas(self, source: Union[str, Path]) -> List[Document]:
        """pandas를 사용한 CSV 로딩"""
        
        documents = []
        
        df = pd.read_csv(source, delimiter=self.delimiter, encoding=self.encoding)
        
        for idx, row in df.iterrows():
            if self.text_column and self.text_column in df.columns:
                content = str(row[self.text_column])
            elif self.combine_columns:
                content = ' | '.join([f"{col}: {str(val)}" for col, val in row.items()])
            else:
                content = str(row.to_dict())
            
            metadata = {
                "file_type": "csv",
                "row_number": idx + 1,
                "total_rows": len(df),
                "columns": list(df.columns),
                "file_size": os.path.getsize(source),
                "loader": "CSVLoader-pandas"
            }
            
            documents.append(Document(
                content=content,
                metadata=metadata,
                source=str(source),
                doc_id=f"{os.path.basename(source)}-row-{idx + 1}"
            ))
        
        return documents
    
    def _load_with_csv(self, source: Union[str, Path]) -> List[Document]:
        """기본 csv 모듈을 사용한 CSV 로딩"""
        
        documents = []
        
        with open(source, 'r', encoding=self.encoding, newline='') as file:
            reader = csv.DictReader(file, delimiter=self.delimiter)
            
            for idx, row in enumerate(reader, 1):
                if self.text_column and self.text_column in row:
                    content = str(row[self.text_column])
                elif self.combine_columns:
                    content = ' | '.join([f"{col}: {str(val)}" for col, val in row.items()])
                else:
                    content = str(row)
                
                metadata = {
                    "file_type": "csv",
                    "row_number": idx,
                    "columns": list(reader.fieldnames),
                    "file_size": os.path.getsize(source),
                    "loader": "CSVLoader-csv"
                }
                
                documents.append(Document(
                    content=content,
                    metadata=metadata,
                    source=str(source),
                    doc_id=f"{os.path.basename(source)}-row-{idx}"
                ))
        
        return documents
    
    def supports_file_type(self, file_path: str) -> bool:
        """CSV 파일 타입 지원 여부"""
        return Path(file_path).suffix.lower() == '.csv'


class JSONLoader(DocumentLoader):
    """JSON 파일 로더"""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__(config)
        
        self.text_key = self.config.get("text_key", None)
        self.flatten_nested = self.config.get("flatten_nested", True)
    
    def load(self, source: Union[str, Path]) -> List[Document]:
        """JSON 파일 로딩"""
        
        if not self.validate_file(source):
            return []
        
        try:
            with open(source, 'r', encoding=self.encoding) as f:
                data = json.load(f)
            
            if isinstance(data, list):
                return self._load_json_array(data, source)
            else:
                return self._load_json_object(data, source)
            
        except Exception as e:
            logger.error(f"JSON 파일 로딩 오류: {source}, {e}")
            return []
    
    def _load_json_array(self, data: List[Dict], source: Union[str, Path]) -> List[Document]:
        """JSON 배열 로딩"""
        
        documents = []
        
        for idx, item in enumerate(data):
            content = self._extract_content(item)
            
            metadata = {
                "file_type": "json",
                "item_index": idx,
                "total_items": len(data),
                "file_size": os.path.getsize(source),
                "loader": "JSONLoader"
            }
            
            documents.append(Document(
                content=content,
                metadata=metadata,
                source=str(source),
                doc_id=f"{os.path.basename(source)}-item-{idx}"
            ))
        
        return documents
    
    def _load_json_object(self, data: Dict, source: Union[str, Path]) -> List[Document]:
        """JSON 객체 로딩"""
        
        content = self._extract_content(data)
        
        metadata = {
            "file_type": "json",
            "keys": list(data.keys()) if isinstance(data, dict) else [],
            "file_size": os.path.getsize(source),
            "loader": "JSONLoader"
        }
        
        return [Document(
            content=content,
            metadata=metadata,
            source=str(source)
        )]
    
    def _extract_content(self, data: Any) -> str:
        """데이터에서 텍스트 컨텐츠 추출"""
        
        if self.text_key and isinstance(data, dict) and self.text_key in data:
            return str(data[self.text_key])
        
        if self.flatten_nested:
            return self._flatten_to_text(data)
        else:
            return json.dumps(data, ensure_ascii=False, indent=2)
    
    def _flatten_to_text(self, data: Any, prefix: str = "") -> str:
        """중첩된 데이터를 평면 텍스트로 변환"""
        
        if isinstance(data, dict):
            items = []
            for key, value in data.items():
                new_prefix = f"{prefix}.{key}" if prefix else key
                items.append(self._flatten_to_text(value, new_prefix))
            return " | ".join(items)
        
        elif isinstance(data, list):
            items = []
            for idx, item in enumerate(data):
                new_prefix = f"{prefix}[{idx}]" if prefix else f"[{idx}]"
                items.append(self._flatten_to_text(item, new_prefix))
            return " | ".join(items)
        
        else:
            return f"{prefix}: {str(data)}" if prefix else str(data)
    
    def supports_file_type(self, file_path: str) -> bool:
        """JSON 파일 타입 지원 여부"""
        return Path(file_path).suffix.lower() == '.json'


def create_document_loader(
    file_type: str,
    config: Optional[Dict[str, Any]] = None
) -> DocumentLoader:
    """파일 타입에 따른 문서 로더 생성"""
    
    file_type = file_type.lower()
    
    loaders = {
        'text': TextLoader,
        'txt': TextLoader,
        'md': TextLoader,
        'pdf': PDFLoader,
        'docx': DocxLoader,
        'csv': CSVLoader,
        'json': JSONLoader,
    }
    
    if file_type not in loaders:
        raise ValueError(f"지원하지 않는 파일 타입: {file_type}")
    
    loader_class = loaders[file_type]
    
    # PDF와 DOCX는 필수 라이브러리 체크
    if file_type == 'pdf' and not HAS_PDF:
        raise ImportError("PDF 로딩을 위해 PyPDF2와 pdfplumber를 설치해주세요")
    elif file_type == 'docx' and not HAS_DOCX:
        raise ImportError("DOCX 로딩을 위해 python-docx를 설치해주세요")
    
    return loader_class(config)


def auto_detect_loader(file_path: Union[str, Path]) -> DocumentLoader:
    """파일 확장자를 기반으로 자동으로 로더 선택"""
    
    file_path = Path(file_path)
    extension = file_path.suffix.lower()
    
    # 확장자별 로더 매핑
    extension_map = {
        '.txt': TextLoader,
        '.md': TextLoader,
        '.rst': TextLoader,
        '.log': TextLoader,
        '.py': TextLoader,
        '.js': TextLoader,
        '.html': TextLoader,
        '.css': TextLoader,
        '.pdf': PDFLoader,
        '.docx': DocxLoader,
        '.csv': CSVLoader,
        '.json': JSONLoader,
    }
    
    if extension not in extension_map:
        logger.warning(f"지원하지 않는 파일 확장자: {extension}, TextLoader를 사용합니다.")
        return TextLoader()
    
    loader_class = extension_map[extension]
    
    try:
        return loader_class()
    except ImportError as e:
        logger.error(f"로더 생성 실패: {e}")
        logger.info("TextLoader를 대신 사용합니다.")
        return TextLoader()